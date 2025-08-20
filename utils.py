import os
import docx
from PyPDF2 import PdfReader

def split_text_into_chunks(text, max_words=3000):
    """Split text into manageable chunks (~10 pages ≈ 3000 words)."""
    words = text.split()
    for i in range(0, len(words), max_words):
        yield " ".join(words[i:i + max_words])

def read_docx(file):
    doc = docx.Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def read_pdf(file):
    pdf = PdfReader(file)
    full_text = []
    for page in pdf.pages:
        full_text.append(page.extract_text())
    return "\n".join(full_text)

def save_to_docx(sections, filename="formatted_novel.docx"):
    doc = docx.Document()
    for section in sections:
        doc.add_paragraph(section)
        doc.add_page_break()
    doc.save(filename)
    return filename
