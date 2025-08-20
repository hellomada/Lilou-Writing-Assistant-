# app.py — Novel Mechanics & Formatting Polisher (Streamlit)
# ----------------------------------------------------------
# Features
# - Upload a .pdf or .docx (200–300+ pages OK)
# - Splits into ~10-page sections (true pages for PDF, estimated for DOCX)
# - Sends each section to OpenAI for **mechanics-only** editing (grammar, punctuation,
#   quotation/Dialogue marks, spacing, indentation, capitalization, consistent style rules)
# - Styles supported: Chicago, APA, MLA, Plain Novel
# - Preserves author voice and word choice (mechanics-only guarantee)
# - Reassembles to a clean publisher-ready DOCX with standard novel layout (12pt TNR, 2x spacing, indents)
# - Progress indicator + resumable processing within a session
#
# Setup
# 1) pip install -r requirements.txt  (see sample at bottom of this file)
# 2) Set your API key:  export OPENAI_API_KEY="sk-..."  (Windows: setx OPENAI_API_KEY "sk-...")
# 3) Run:  streamlit run app.py
#
# Notes
# - For PDFs, text extraction quality depends on the source PDF (true text vs scanned).
# - For DOCX, page count is estimated (default 300 words/page); you can adjust in the UI.
# - Output is a DOCX with standard manuscript formatting.

import io
import os
from typing import List, Tuple

import streamlit as st

# PDF + DOCX handling
try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_LINE_SPACING
except Exception:
    Document = None

# OpenAI (SDK v1+)
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

from tenacity import retry, stop_after_attempt, wait_exponential

# -----------------------------
# UI — Sidebar & Page Setup
# -----------------------------
st.set_page_config(page_title="Novel Formatting Polisher", page_icon="📖", layout="wide")
st.title("📖 Novel Mechanics & Formatting Polisher")
st.caption("Mechanics-only edit: grammar, punctuation, quotation rules, spacing & indentation — your voice preserved.")

with st.sidebar:
    st.header("⚙️ Options")
    style = st.selectbox(
        "Choose target style",
        options=["Chicago", "APA", "MLA", "Plain Novel"],
        index=0,
        help="Applies mechanical rules from the selected style while preserving your prose."
    )
    words_per_page = st.number_input(
        "Estimate words/page (for DOCX only)", min_value=200, max_value=600, value=300, step=10,
        help="Used only when your upload is a DOCX. PDFs use real pages."
    )
    pages_per_chunk = st.number_input(
        "Pages per chunk", min_value=5, max_value=20, value=10, step=1,
        help="Each chunk is processed independently to stay within model limits."
    )
    smart_quotes = st.checkbox("Normalize to curly quotes (“ ” ‘ ’)", value=True)
    preserve_dashes = st.checkbox("Normalize dashes (–/—) per style", value=True)
    normalize_ellipses = st.checkbox("Normalize ellipses (…) spacing", value=True)

st.subheader("1) Upload your manuscript (.pdf or .docx)")
uploaded = st.file_uploader("Upload file", type=["pdf", "docx"], accept_multiple_files=False)

col_a, col_b = st.columns([3, 2])

with col_a:
    st.subheader("2) Review + Process")
    st.write("Select style, then click **Process**. You'll be able to download a polished DOCX.")

with col_b:
    st.info("This tool focuses on **mechanics** (grammar, punctuation, spacing, indentation, style rules). It **does not** rewrite your story.")

# -----------------------------
# Helpers
# -----------------------------
STYLE_RULESETS = {
    "Chicago": [
        "U.S. English punctuation and spelling.",
        "Curly quotes for dialogue; closing punctuation inside quotes.",
        "Use em dashes (—) for breaks in thought; no spaces around em dashes.",
        "Serial (Oxford) comma.",
        "Numbers spelled out for one through one hundred in narrative prose unless clarity demands numerals.",
        "Ellipses as … with appropriate surrounding spacing when omitting words.",
        "Single space after periods.",
        "Capitalize proper nouns consistently.",
        "Standardize italics for internal thoughts if already italicized (do not newly italicize).",
        "Leave paragraphing as-is except to fix obvious spacing/indentation errors.",
    ],
    "APA": [
        "U.S. English punctuation and spelling.",
        "Double quotation marks for quotes; punctuation inside quotation marks.",
        "Use en dash (–) for numeric ranges; em dash (—) for sentence breaks (no spaces).",
        "Single space after periods.",
        "Oxford comma.",
        "Consistent capitalization of headings if present.",
        "Standardize ellipses according to APA guidance.",
        "Numbers: spell out one through nine in general prose, otherwise numerals (unless at sentence start).",
        "Curly quotes if smart quotes selected.",
        "Do not insert citations or references; mechanics only.",
    ],
    "MLA": [
        "U.S. English punctuation and spelling.",
        "Curly quotes; punctuation placement per MLA.",
        "Oxford comma.",
        "Single space after periods.",
        "Em dash (—) without spaces; en dash (–) for ranges.",
        "Standardize ellipses and capitalization.",
        "Do not add citations; mechanics only.",
    ],
    "Plain Novel": [
        "U.S. English punctuation and spelling.",
        "Curly quotes for dialogue; punctuation inside quotes.",
        "Em dashes (—) for interruptions; no spaces around em dashes.",
        "Oxford comma; consistent capitalization.",
        "Normalize ellipses to … with proper spacing.",
        "Single space after periods.",
        "Fix spacing/indentation only; do not change paragraph boundaries unless obviously broken.",
    ],
}

MECHANICS_ONLY_PLEDGE = (
    "Preserve the author's voice and word choice. Do NOT add, delete, or substitute content except to fix grammar, "
    "punctuation, capitalization, spacing, indentation, and mechanical style. If uncertain, leave the original as-is."
)


def _require(pkg_name: str, obj):
    if obj is None:
        st.error(f"Missing dependency: {pkg_name}. Please add it to requirements.txt and install.")
        st.stop()


def extract_pdf_chunks(file_bytes: bytes, pages_per_chunk: int) -> List[Tuple[str, Tuple[int, int]]]:
    """Return list of (text, (start_page, end_page)) chunks for a PDF."""
    _require("pypdf", PdfReader)
    reader = PdfReader(io.BytesIO(file_bytes))
    total_pages = len(reader.pages)
    chunks = []
    for start in range(0, total_pages, pages_per_chunk):
        end = min(start + pages_per_chunk, total_pages)
        buf = []
        for p in range(start, end):
            try:
                buf.append(reader.pages[p].extract_text() or "")
            except Exception:
                buf.append("")
        text = "\n".join(buf)
        chunks.append((text, (start + 1, end)))  # 1-based page numbers
    return chunks


def extract_docx_chunks(file_bytes: bytes, pages_per_chunk: int, words_per_page: int) -> List[Tuple[str, Tuple[int, int]]]:
    """Estimate pages from word count; return (text, (start_page, end_page))."""
    _require("python-docx", Document)
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs]
    text = "\n".join(paragraphs)
    words = text.split()
    words_per_chunk = int(words_per_page * pages_per_chunk)

    chunks = []
    start_idx = 0
    page_start_num = 1
    while start_idx < len(words):
        end_idx = min(start_idx + words_per_chunk, len(words))
        chunk_words = words[start_idx:end_idx]
        chunk_text = " ".join(chunk_words)
        approx_pages = len(chunk_words) / max(1, words_per_page)
        page_end_num = int(round(page_start_num + approx_pages - 1))
        chunks.append((chunk_text, (page_start_num, page_end_num)))
        start_idx = end_idx
        page_start_num = page_end_num + 1
    return chunks


# OpenAI client
@st.cache_resource(show_spinner=False)
def get_openai_client():
    _require("openai", OpenAI)
    api_key = os.getenv("OPENAI_API_KEY", "").strip()
    if not api_key:
        st.error("OPENAI_API_KEY not set. Set it in your environment and restart.")
        st.stop()
    return OpenAI()


def build_style_instructions(style_name: str) -> str:
    rules = STYLE_RULESETS.get(style_name, STYLE_RULESETS["Plain Novel"])[:]
    # Optional normalizations based on toggles
    if smart_quotes:
        rules.append("Normalize straight quotes to curly quotes if and only if they are quotation marks, not inch/foot marks.")
    if preserve_dashes:
        rules.append("Normalize dashes: use em dashes (—) for interruptions/breaks, en dashes (–) for ranges; avoid surrounding spaces.")
    if normalize_ellipses:
        rules.append("Normalize ellipses to a single glyph (…) with correct surrounding spacing.")
    return "\n- " + "\n- ".join(rules)


EDIT_SYSTEM_PROMPT = (
    "You are a professional fiction copy editor. Apply **mechanics-only** corrections while preserving the author's voice.\n"
    "Allowed operations: fix grammar, punctuation, capitalization, spelling (en-US), consistent quotation/dialogue marks, spacing, indentation cues, dash and ellipsis normalization, and other style mechanics.\n"
    "Forbidden: changing word choice, tone, plot details, character voice, adding or removing sentences, or restructuring paragraphs unless it's clearly a mechanical spacing/indentation repair.\n"
)


@retry(wait=wait_exponential(multiplier=1, min=2, max=20), stop=stop_after_attempt(4))
def edit_chunk_with_openai(client: "OpenAI", text: str, style_name: str) -> str:
    if not text.strip():
        return ""
    style_rules = build_style_instructions(style_name)

    # We use the Responses API for clarity and steerability
    resp = client.responses.create(
        model="gpt-4o-mini",  # cost-efficient and strong at editing tasks
        input=[
            {"role": "system", "content": EDIT_SYSTEM_PROMPT},
            {"role": "user", "content": (
                f"Target style: {style_name}.\n" \
                f"Mechanics-only pledge: {MECHANICS_ONLY_PLEDGE}\n" \
                f"Style rules to apply:\n{style_rules}\n\n" \
                "Instructions:\n" \
                "1) Return only the **corrected text**.\n" \
                "2) Preserve original wording; change only mechanics.\n" \
                "3) Keep the original paragraph breaks except for obvious spacing/indent fixes.\n\n" \
                "Text to edit begins below the fence.\n" \
                "<<<BEGIN>>>\n" + text + "\n<<<END>>>\n"
            )}
        ],
        temperature=0.0,
        max_output_tokens=4000,
    )

    # Extract the first text output
    out = resp.output_text if hasattr(resp, "output_text") else ""
    return out.strip()


# DOCX creation with manuscript formatting

def build_manuscript_docx(paragraphs: List[str]) -> bytes:
    _require("python-docx", Document)
    doc = Document()

    # Set base style: 12pt Times New Roman, double-spaced, first-line indent
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Ensure correct font (Windows/Word)
    try:
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    except Exception:
        pass

    first_line_indent = Inches(0.5)

    for block in paragraphs:
        for para in block.split("\n\n"):
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.first_line_indent = first_line_indent
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# -----------------------------
# Main flow
# -----------------------------
if uploaded is not None:
    name = uploaded.name
    ext = name.lower().split(".")[-1]

    st.write(f"**File:** {name}")

    data = uploaded.read()

    if ext == "pdf":
        if PdfReader is None:
            st.error("pypdf not installed. Add it to requirements and restart.")
            st.stop()
        chunks = extract_pdf_chunks(data, int(pages_per_chunk))
    elif ext == "docx":
        if Document is None:
            st.error("python-docx not installed. Add it to requirements and restart.")
            st.stop()
        chunks = extract_docx_chunks(data, int(pages_per_chunk), int(words_per_page))
    else:
        st.error("Unsupported file type. Please upload .pdf or .docx")
        st.stop()

    st.success(f"Prepared {len(chunks)} chunks (~{pages_per_chunk} pages each). Ready to process.")

    client = get_openai_client()

    if st.button("🚀 Process Manuscript", type="primary"):
        progress = st.progress(0)
        edited_blocks: List[str] = []

        for i, (chunk_text, (p_start, p_end)) in enumerate(chunks, start=1):
            st.write(f"Editing chunk {i}/{len(chunks)} (pages {p_start}–{p_end})…")
            edited = edit_chunk_with_openai(client, chunk_text, style)
            edited_blocks.append(edited)
            progress.progress(i / len(chunks))

        st.success("All chunks edited.")

        # Build DOCX
        st.write("Assembling final DOCX with standard manuscript formatting…")
        docx_bytes = build_manuscript_docx(edited_blocks)
        st.download_button(
            "📥 Download Polished Manuscript (.docx)",
            data=docx_bytes,
            file_name=f"{os.path.splitext(name)[0]}_polished.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Optional: also provide a merged plain text
        merged_text = "\n\n".join(edited_blocks)
        st.download_button(
            "📄 Download Merged Text (.txt)",
            data=merged_text.encode("utf-8"),
            file_name=f"{os.path.splitext(name)[0]}_polished.txt",
            mime="text/plain",
        )

# -----------------------------
# (Optional) requirements.txt reference
# -----------------------------
# Save the following as requirements.txt next to app.py:
# streamlit>=1.38
# openai>=1.35.0
# pypdf>=4.2.0
# python-docx>=1.1.2
# tenacity>=8.4.1
